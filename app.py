import io
import re
from typing import Optional, Tuple, List

import streamlit as st
import pandas as pd

# ---------- Helpers ----------
def clean_phone(s: str) -> Optional[str]:
    """Очищает и валидирует номер телефона"""
    # Убираем все кроме цифр и +
    digits = re.sub(r"[^\d+]", "", s)
    
    # Убираем множественные +
    digits = re.sub(r'\++', '+', digits)
    
    # Убираем + не в начале
    if '+' in digits[1:]:
        digits = digits[0] + digits[1:].replace('+', '')
    
    # Получаем только цифры для проверки длины
    core = re.sub(r"\D", "", digits)
    
    # Украинские номера: 10 цифр (0XXXXXXXXX) или 12 цифр (+380XXXXXXXXX)
    if len(core) == 10 and core.startswith('0'):
        return core
    elif len(core) == 12 and core.startswith('380'):
        return '+' + core
    elif len(core) == 9:  # Без первого 0
        return '0' + core
    elif 9 <= len(core) <= 13:  # Другие форматы
        return digits
    
    return None


def find_all_phones(text: str) -> List[str]:
    """Находит все телефоны в тексте"""
    candidates = []
    
    # Убираем лишние пробелы
    text = re.sub(r'\s+', ' ', text)
    
    # Множество паттернов для разных форматов
    patterns = [
        # +380XXXXXXXXX (с возможными разделителями)
        r'\+\s?380[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d',
        # 380XXXXXXXXX (без +)
        r'\b380[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d',
        # 0XXXXXXXXX (украинский формат)
        r'\b0[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d',
        # Любые 10 цифр подряд (с разделителями)
        r'\b\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d',
        # Просто 10 цифр подряд
        r'\b\d{10}\b',
        # 9 цифр (может не хватать первого 0)
        r'\b\d{9}\b',
    ]
    
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            phone_raw = match.group(0)
            phone = clean_phone(phone_raw)
            if phone and phone not in candidates:
                # Проверяем что это не дата или ID
                digits_only = re.sub(r'\D', '', phone)
                
                # Пропускаем если похоже на дату (начинается с 19, 20, 01-31)
                if len(digits_only) == 8 and (digits_only.startswith('19') or digits_only.startswith('20')):
                    continue
                
                # Пропускаем если все цифры одинаковые
                if len(set(digits_only)) == 1:
                    continue
                
                # Пропускаем ID документов (слишком длинные)
                if len(digits_only) > 13:
                    continue
                
                candidates.append(phone)
    
    return candidates


def best_phone(text: str) -> Optional[str]:
    """Возвращает первый найденный телефон"""
    phones = find_all_phones(text)
    
    # Приоритет: сначала номера с +380 или начинающиеся с 0
    priority_phones = [p for p in phones if p.startswith('+380') or p.startswith('0')]
    if priority_phones:
        return priority_phones[0]
    
    return phones[0] if phones else None


def read_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
    except Exception:
        return ""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def read_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    f = io.BytesIO(file_bytes)
    doc = Document(f)
    parts = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def read_file(uploaded_file) -> str:
    data = uploaded_file.read()
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return read_pdf(data)
    if name.endswith(".docx"):
        return read_docx(data)
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""


# ---------- Extraction from text ----------
def extract_fio_from_text(text: str) -> Optional[str]:
    """Извлекает ФИО из документа резюме"""
    if not text:
        return None
    
    # Убираем лишние пробелы и разбиваем на строки
    text = re.sub(r'\s+', ' ', text)
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    
    # Ищем в первых 50 строках
    search_lines = lines[:50]
    
    # Паттерны для поиска ФИО - ОЧЕНЬ гибкие
    fio_patterns = [
        # 3 слова с заглавной (Фамилия Имя Отчество)
        r'([А-ЯЁІЇЄҐA-Z][а-яёіїєґa-z]{2,})\s+([А-ЯЁІЇЄҐA-Z][а-яёіїєґa-z]{2,})\s+([А-ЯЁІЇЄҐA-Z][а-яёіїєґa-z]{2,})',
        # 2 слова с заглавной (Фамилия Имя) - минимум 2 буквы
        r'([А-ЯЁІЇЄҐA-Z][а-яёіїєґa-z]{1,})\s+([А-ЯЁІЇЄҐA-Z][а-яёіїєґa-z]{1,})',
    ]
    
    # Минимальный список стоп-слов - только явные служебные слова
    stop_words = [
        'резюме', 'curriculum', 'vitae', 
        'email', 'www', 'http', 'https',
        'розглядає', 'рассматривает',
        'январ', 'феврал', 'март', 'апрел', 'май', 'июн', 'июл', 'август', 'сентябр', 'октябр', 'ноябр', 'декабр',
        'січня', 'лютого', 'березня', 'квітня', 'травня', 'червня', 'липня', 'серпня', 'вересня', 'жовтня', 'листопада', 'грудня',
        'january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december',
        'понедельник', 'вторник', 'среда', 'четверг', 'пятница', 'суббота', 'воскресенье',
        'понеділок', 'вівторок', 'середа', 'четвер', "п'ятниця", 'субота', 'неділя',
        'university', 'університет', 'інститут', 'institute', 'академія', 'academy', 'школа', 'school',
        'освіта', 'образование', 'education', 'досвід', 'опыт', 'experience'
    ]
    
    # Слова должностей - НЕ стоп-слова, ФИО может быть рядом
    position_words = [
        'менеджер', 'manager', 'адміністратор', 'administrator', 'продавець', 'продавец',
        'консультант', 'спеціаліст', 'специалист', 'specialist', 'помічник', 'помощник',
        'оператор', 'operator', 'директор', 'director', 'координатор', 'рекрутер', 'recruiter'
    ]
    
    candidates = []
    
    for idx, line in enumerate(search_lines):
        line_lower = line.lower()
        
        # Пропускаем только явно нерелевантные строки
        # Email и URL
        if re.search(r'@|https?://|www\.', line):
            continue
        
        # Строки с доменами
        if re.search(r'\.(com|ru|ua|org|net|gov)', line):
            continue
        
        # Даты в формате дд.мм.гггг или просто год
        if re.search(r'\b\d{2}\.\d{2}\.\d{4}\b|\b(199|200|201|202)\d\b', line):
            continue
        
        # Строки где больше 10 цифр (скорее всего номера/коды)
        digit_count = sum(c.isdigit() for c in line)
        if digit_count > 10:
            continue
        
        # Пропускаем если вся строка - это только стоп-слова
        if any(line_lower.startswith(word) for word in stop_words):
            if not any(pw in line_lower for pw in position_words):
                continue
        
        # Ищем ФИО по паттернам
        for pattern in fio_patterns:
            matches = list(re.finditer(pattern, line))
            for match in matches:
                fio = ' '.join(match.groups())
                words = fio.split()
                
                # Минимальные проверки
                if len(words) < 2:
                    continue
                
                # Проверяем что слова не стоп-слова
                if any(word.lower() in stop_words for word in words):
                    continue
                
                # Пропускаем если все слова - названия должностей
                if all(word.lower() in position_words for word in words):
                    continue
                
                # Считаем приоритет
                priority = 100 - idx  # Позиция в документе
                
                # Бонусы
                if len(words) == 3:  # ФИО с отчеством
                    priority += 60
                elif len(words) == 2:
                    # Смотрим на длину слов
                    avg_len = sum(len(w) for w in words) / len(words)
                    if avg_len >= 5:  # Полные слова
                        priority += 40
                    elif avg_len >= 3:  # Средние/сокращенные
                        priority += 25
                    else:
                        priority += 10
                
                # Огромный бонус если в первых 3 строках
                if idx < 3:
                    priority += 100
                elif idx < 10:
                    priority += 50
                
                # Бонус если строка короткая (вероятно только ФИО)
                if len(line) < 60:
                    priority += 30
                
                # Штраф если есть должность в строке
                if any(pw in line_lower for pw in position_words):
                    priority -= 10
                
                candidates.append((priority, fio, idx, line))
    
    # Выбираем лучшего кандидата
    if candidates:
        candidates.sort(key=lambda x: x[0], reverse=True)
        # Возвращаем топ кандидата
        best = candidates[0]
        return best[1]
    
    return None


def extract_position_from_text(text: str) -> Optional[str]:
    """Извлекает желаемую должность из резюме"""
    if not text:
        return None
    
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    
    # Паттерны для явного указания должности
    position_patterns = [
        r'(?i)(?:желаемая|бажана)\s+(?:должность|посада)[:\s\-—]*(.+?)(?:\n|$)',
        r'(?i)(?:должность|посада)[:\s\-—]+(.+?)(?:\n|$)',
        r'(?i)(?:вакансия|вакансія)[:\s\-—]+(.+?)(?:\n|$)',
        r'(?i)(?:розглядає|рассматривает)\s+(?:посади|должности)[:\s\-—]*(.+?)(?:\n|$)',
        r'(?i)(?:позиция|позиція)[:\s\-—]+(.+?)(?:\n|$)',
        r'(?i)position[:\s\-—]+(.+?)(?:\n|$)',
        r'(?i)objective[:\s\-—]+(.+?)(?:\n|$)',
        r'(?i)(?:цель|ціль)[:\s\-—]+(.+?)(?:\n|$)',
        r'(?i)(?:ищу|шукаю)\s+(?:работу|роботу)[:\s\-—]*(.+?)(?:\n|$)',
    ]
    
    # Ищем по всему тексту
    for pattern in position_patterns:
        match = re.search(pattern, text[:4000])
        if match:
            position = match.group(1).strip()
            # Очистка
            position = re.sub(r'[_\-—\.]+$', '', position).strip()
            position = position.split('\n')[0].strip()
            
            # Убираем ID и номера в конце
            position = re.sub(r'\s+\d{6,}$', '', position)
            
            if 3 <= len(position) <= 200:
                return position
    
    # Если не нашли явное указание, ищем в первых 40 строках
    position_keywords_uk = [
        'менеджер', 'адміністратор', 'продавець', 'консультант', 'спеціаліст',
        'помічник', 'оператор', 'координатор', 'асистент', 'керівник',
        'бариста', 'офіс-менеджер', 'секретар', 'рекрутер'
    ]
    
    position_keywords_ru = [
        'менеджер', 'администратор', 'продавец', 'консультант', 'специалист',
        'помощник', 'оператор', 'координатор', 'ассистент', 'руководитель',
        'бариста', 'офис-менеджер', 'секретарь', 'рекрутер'
    ]
    
    position_keywords_en = [
        'manager', 'administrator', 'seller', 'consultant', 'specialist',
        'assistant', 'operator', 'coordinator', 'director', 'supervisor',
        'barista', 'recruiter', 'designer', 'developer', 'engineer'
    ]
    
    all_keywords = position_keywords_uk + position_keywords_ru + position_keywords_en
    
    for i, line in enumerate(lines[:40]):
        line_lower = line.lower()
        
        # Пропускаем слишком длинные
        if len(line) > 150:
            continue
        
        # Пропускаем с email, url, датами
        if re.search(r'(@|https?://|www\.|\d{4}|\.com|\.ua|\.ru)', line):
            continue
        
        # Ищем ключевые слова должностей
        if any(keyword in line_lower for keyword in all_keywords):
            # Убираем ФИО из строки если есть
            cleaned = line
            
            # Убираем ID номера
            cleaned = re.sub(r'\s+\d{6,}$', '', cleaned)
            
            # Если в строке есть слова с заглавной буквы и должность, пытаемся разделить
            words = cleaned.split()
            position_words = []
            
            for word in words:
                word_lower = word.lower()
                # Если слово похоже на должность
                if any(kw in word_lower for kw in all_keywords):
                    position_words.append(word)
                # Или если слово уже добавлено и текущее может быть частью должности
                elif position_words and not re.match(r'^[А-ЯЁA-Z][а-яёa-z]+$', word):
                    position_words.append(word)
                elif position_words:
                    position_words.append(word)
            
            if position_words:
                position = ' '.join(position_words).strip()
                position = re.sub(r'[,;:.]+$', '', position)
                
                # Финальная проверка длины
                if 5 <= len(position) <= 150:
                    return position
            
            # Если не получилось разделить, берем всю строку
            if 5 <= len(cleaned) <= 100:
                return cleaned
    
    return None


def extract_position_from_filename(filename: str) -> str:
    """Извлекает должность из имени файла"""
    name_clean = re.sub(r"\.[^.]+$", "", filename)
    
    # Если только цифры - возвращаем "—"
    if re.match(r'^\d+$', name_clean):
        return "—"
    
    try:
        from urllib.parse import unquote
        name_clean = unquote(name_clean)
    except:
        pass
    
    # Очистка
    for pattern in [r'(?i)workua', r'(?i)work\.ua', r'(?i)резюме', r'(?i)resume', r'(?i)\bcv\b']:
        name_clean = re.sub(pattern, ' ', name_clean)
    
    name_clean = re.sub(r'\b\d{10,}\b', ' ', name_clean)
    name_clean = re.sub(r'\d{2,4}[-./]\d{1,2}[-./]\d{2,4}', ' ', name_clean)
    name_clean = re.sub(r"[_\-—,\.]+", " ", name_clean)
    name_clean = re.sub(r'\s+', ' ', name_clean).strip()
    
    return name_clean if name_clean else "—"


# ---------- UI ----------
st.set_page_config(page_title="CV Extractor", page_icon="📄", layout="centered")
st.title("📄 Екстрактор даних з резюме")
st.caption("Автоматичне витягування ПІБ, посади та телефону")

debug_mode = st.sidebar.checkbox("🔍 Режим відладки", value=False)
show_stats = st.sidebar.checkbox("📊 Показати статистику", value=True)

uploaded = st.file_uploader(
    "Завантажте файли резюме (PDF/DOCX)", 
    accept_multiple_files=True,
    type=["pdf", "docx"],
)

if uploaded:
    rows = []
    debug_info = []
    stats = {'fio_found': 0, 'position_found': 0, 'phone_found': 0, 'total': len(uploaded)}
    
    with st.spinner(f'Обробка {len(uploaded)} файлів...'):
        for uf in uploaded:
            text = read_file(uf)
            
            # Извлечение данных
            fio = extract_fio_from_text(text) or "—"
            position_text = extract_position_from_text(text)
            position_filename = extract_position_from_filename(uf.name)
            position = position_text if position_text else position_filename
            phone = best_phone(text) or "—"
            
            # Статистика
            if fio != "—":
                stats['fio_found'] += 1
            if position != "—":
                stats['position_found'] += 1
            if phone != "—":
                stats['phone_found'] += 1
            
            rows.append({
                "Файл": uf.name,
                "ПІБ": fio,
                "Бажана посада": position,
                "Телефон": phone,
            })
            
            if debug_mode:
                all_phones = find_all_phones(text)
                
                # Находим всех кандидатов на ФИО для отладки
                fio_candidates = []
                lines_preview = [l.strip() for l in text.split('\n') if l.strip()][:50]
                
                for idx, line in enumerate(lines_preview):
                    # Простой поиск всех слов с заглавной
                    pattern = r'([А-ЯЁІЇЄҐA-Z][а-яёіїєґa-z]{1,})\s+([А-ЯЁІЇЄҐA-Z][а-яёіїєґa-z]{1,})'
                    matches = re.finditer(pattern, line)
                    for match in matches:
                        candidate = ' '.join(match.groups())
                        fio_candidates.append(f"Строка {idx+1}: {candidate} | Вся строка: {line[:80]}")
                
                # Анализ текста на наличие последовательностей цифр
                digit_sequences = []
                # Ищем любые последовательности из 9-12 цифр
                for match in re.finditer(r'\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d[\s\-\(\)\.]*\d', text[:2000]):
                    seq = match.group(0)
                    digit_sequences.append(seq)
                
                debug_info.append({
                    "Файл": uf.name,
                    "ПІБ (результат)": fio,
                    "Посада (з тексту)": position_text or "—",
                    "Посада (з назви)": position_filename,
                    "Всі телефони": ', '.join(all_phones) if all_phones else "—",
                    "Кандидати на ПІБ": '\n'.join(fio_candidates[:15]) if fio_candidates else "Не знайдено жодного",
                    "Послідовності цифр": ', '.join(digit_sequences[:10]) if digit_sequences else "Не знайдено",
                    "Текст (початок)": text[:800] if text else "❌ Текст не отримано"
                })
    
    df = pd.DataFrame(rows)
    
    # Показываем статистику
    if show_stats:
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("📁 Всього файлів", stats['total'])
        col2.metric("👤 ПІБ знайдено", f"{stats['fio_found']}/{stats['total']}")
        col3.metric("💼 Посад знайдено", f"{stats['position_found']}/{stats['total']}")
        col4.metric("📞 Телефонів знайдено", f"{stats['phone_found']}/{stats['total']}")
    
    st.dataframe(df, use_container_width=True, height=400)
    
    # Отладка
    if debug_mode and debug_info:
        st.subheader("🔍 Детальна інформація")
        st.caption("Тут показано що саме знайдено в кожному файлі та ВСІ можливі кандидати на ПІБ")
        
        for info in debug_info:
            status = "✅" if info['ПІБ (результат)'] != "—" else "❌"
            with st.expander(f"{status} {info['Файл'][:60]}"):
                
                st.write("### 🎯 Результат")
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**ПІБ (обрано):**")
                    if info['ПІБ (результат)'] != "—":
                        st.success(info['ПІБ (результат)'])
                    else:
                        st.error("Не знайдено")
                    
                    st.write("**Телефон (обрано):**")
                    if info['Всі телефони'] != "—":
                        st.success(info['Всі телефони'].split(',')[0])
                    else:
                        st.error("Не знайдено")
                
                with col2:
                    st.write("**Посада з тексту:**")
                    if info['Посада (з тексту)'] != "—":
                        st.info(info['Посада (з тексту)'])
                    else:
                        st.warning("Не знайдено")
                    
                    st.write("**Всі знайдені телефони:**")
                    if info['Всі телефони'] != "—":
                        st.info(info['Всі телефони'])
                    else:
                        st.error("Жодного телефону не знайдено")
                
                st.write("### 👥 Всі знайдені кандидати на ПІБ")
                st.text(info['Кандидати на ПІБ'])
                
                st.write("### 📞 Аналіз пошуку телефонів")
                st.text(f"Знайдені послідовності цифр: {info['Послідовності цифр']}")
                
                st.write("### 📄 Початок тексту документа (перші 800 символів)")
                st.code(info['Текст (початок)'], language='text')
    
    # CSV экспорт
    csv = df.to_csv(index=False).encode("utf-8-sig")
    st.download_button(
        label="📥 Завантажити CSV",
        data=csv,
        file_name="cv_extract.csv",
        mime="text/csv",
    )
else:
    st.info("👆 Завантажте файли резюме для обробки")

st.markdown("""
---
### 📋 Як працює:
- **ПІБ**: шукається в перших 35 рядках документа, підтримує повні та скорочені імена
- **Посада**: шукає за ключовими словами "Розглядає посади:", або автоматично визначає
- **Телефон**: знаходить українські номери (0XXXXXXXXX або +380XXXXXXXXX)

### 💡 Поради:
- Увімкніть "Режим відладки" щоб побачити, що саме знайдено в кожному файлі
- Якщо ПІБ не знайдено, перевірте що воно є в перших рядках документа
""")